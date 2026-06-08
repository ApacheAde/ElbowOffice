#!/usr/bin/env python3
"""
ElbowOffice - A simple, powerful word processor written in Python
Inspired by OpenOffice Writer / LibreOffice Writer / Microsoft Word

A complete WYSIWYG rich-text word processor using PyQt6.

Features:
- Multiple document tabs
- Full rich text formatting (fonts, sizes, colors, bold/italic/underline/strike)
- Paragraph alignment, lists (bullets & numbering), indentation
- Insert images and tables
- Open/Save as HTML (excellent fidelity), Plain Text, and optional DOCX
- Export to PDF and direct printing
- Word/character count, cursor position
- Find & Replace
- Unsaved changes protection
- Keyboard shortcuts throughout

Installation (Windows / macOS / Linux):
    pip install PyQt6

Optional (for .docx support):
    pip install python-docx

Run:
    python elbowoffice.py

"""

import sys
import os

# ====================== FRIENDLY DEPENDENCY CHECK ======================
try:
    from PyQt6.QtWidgets import (
        QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
        QTabWidget, QTextEdit, QToolBar, QStatusBar, QMenuBar, QFileDialog,
        QMessageBox, QInputDialog, QFontComboBox, QComboBox, QColorDialog,
        QDialog, QDialogButtonBox, QLabel, QSpinBox, QLineEdit, QPushButton,
        QGroupBox, QFormLayout
    )
    from PyQt6.QtGui import (
        QAction, QIcon, QFont, QTextCharFormat, QTextBlockFormat,
        QTextCursor, QColor, QTextListFormat, QTextImageFormat,
        QPageLayout, QTextDocument, QKeySequence, QTextDocumentWriter
    )
    from PyQt6.QtCore import Qt, QSize, QSettings, QUrl
    from PyQt6.QtPrintSupport import QPrinter, QPrintDialog, QPrintPreviewDialog
except ImportError:
    print("=" * 70)
    print("ElbowOffice requires PyQt6")
    print("=" * 70)
    print()
    print("Please install it with:")
    print("    pip install PyQt6")
    print()
    print("On Windows you can also use:")
    print("    winget install Python.Python.3.12")
    print("    pip install PyQt6")
    print()
    print("For .docx support (optional):")
    print("    pip install python-docx")
    print()
    input("Press Enter to exit...")
    sys.exit(1)

# Optional DOCX support
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ====================== HELPER FUNCTIONS ======================

def get_desktop_path():
    """Return the user's Desktop path cross-platform."""
    if sys.platform == "win32":
        return os.path.join(os.path.expanduser("~"), "Desktop")
    return os.path.expanduser("~/Desktop")


def resource_path(relative_path):
    """For future bundling (PyInstaller etc)."""
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


# ====================== CUSTOM WIDGETS ======================

class FindReplaceDialog(QDialog):
    """Simple Find & Replace dialog."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Find & Replace")
        self.setModal(False)
        self.resize(420, 160)

        layout = QVBoxLayout(self)

        # Find row
        find_layout = QHBoxLayout()
        find_layout.addWidget(QLabel("Find:"))
        self.find_edit = QLineEdit()
        self.find_edit.setPlaceholderText("Text to find...")
        find_layout.addWidget(self.find_edit)
        layout.addLayout(find_layout)

        # Replace row
        replace_layout = QHBoxLayout()
        replace_layout.addWidget(QLabel("Replace:"))
        self.replace_edit = QLineEdit()
        self.replace_edit.setPlaceholderText("Replacement text...")
        replace_layout.addWidget(self.replace_edit)
        layout.addLayout(replace_layout)

        # Options
        options_layout = QHBoxLayout()
        self.case_sensitive = QPushButton("Case Sensitive")
        self.case_sensitive.setCheckable(True)
        self.whole_word = QPushButton("Whole Words")
        self.whole_word.setCheckable(True)
        options_layout.addWidget(self.case_sensitive)
        options_layout.addWidget(self.whole_word)
        options_layout.addStretch()
        layout.addLayout(options_layout)

        # Buttons
        btn_layout = QHBoxLayout()
        self.find_next_btn = QPushButton("Find Next")
        self.replace_btn = QPushButton("Replace")
        self.replace_all_btn = QPushButton("Replace All")
        self.close_btn = QPushButton("Close")
        btn_layout.addWidget(self.find_next_btn)
        btn_layout.addWidget(self.replace_btn)
        btn_layout.addWidget(self.replace_all_btn)
        btn_layout.addStretch()
        btn_layout.addWidget(self.close_btn)
        layout.addLayout(btn_layout)

        self.close_btn.clicked.connect(self.close)

        self.setLayout(layout)


class DocumentTab(QWidget):
    """
    Represents one open document.
    Contains the QTextEdit and tracks file path + modification state.
    """
    def __init__(self, parent=None):
        super().__init__(parent)
        self.file_path = None
        self.original_title = "Untitled"

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)

        self.editor = QTextEdit()
        self.editor.setAcceptRichText(True)
        self.editor.setAutoFormatting(QTextEdit.AutoFormattingFlag.AutoAll)
        self.editor.setFont(QFont("Calibri", 11))
        self.editor.setTabStopDistance(40)  # Nice tab width

        # Default document margins (like Word)
        doc = self.editor.document()
        doc.setDocumentMargin(25)
        doc.setModified(False)

        layout.addWidget(self.editor)
        self.setLayout(layout)

        # Connect modification signal
        self.editor.document().modificationChanged.connect(self._on_modification_changed)

    def _on_modification_changed(self, changed):
        # This will be handled by the main window via signals
        pass

    def is_modified(self):
        return self.editor.document().isModified()

    def set_modified(self, value: bool):
        self.editor.document().setModified(value)

    def get_title(self):
        if self.file_path:
            return os.path.basename(self.file_path)
        return self.original_title

    def word_count(self):
        text = self.editor.toPlainText()
        words = [w for w in text.split() if w.strip()]
        return len(words)

    def char_count(self):
        return len(self.editor.toPlainText())


# ====================== MAIN APPLICATION ======================

class ElbowOffice(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("ElbowOffice - Word Processor")
        self.resize(1100, 750)
        self.setMinimumSize(800, 500)

        self.settings = QSettings("ElbowOffice", "ElbowOfficeApp")
        self.recent_files = self.settings.value("recent_files", [], type=list)

        self.current_find_dialog = None

        self._setup_ui()
        self._create_actions()
        self._create_menus()
        self._create_toolbars()
        self._create_statusbar()

        # Open with one blank document
        self.new_document()

        # Restore window geometry
        if self.settings.value("geometry"):
            self.restoreGeometry(self.settings.value("geometry"))

        self.update_recent_files_menu()
        self.statusBar().showMessage("Ready", 2000)

    # ---------------- UI SETUP ----------------
    def _setup_ui(self):
        # Central tab widget for multiple documents
        self.tabs = QTabWidget()
        self.tabs.setTabsClosable(True)
        self.tabs.setMovable(True)
        self.tabs.setDocumentMode(True)
        self.tabs.tabCloseRequested.connect(self.close_tab)
        self.tabs.currentChanged.connect(self.on_tab_changed)

        self.setCentralWidget(self.tabs)

    def _create_actions(self):
        # File actions
        self.new_act = QAction("&New", self, shortcut=QKeySequence.StandardKey.New,
                               statusTip="Create a new document", triggered=self.new_document)
        self.open_act = QAction("&Open...", self, shortcut=QKeySequence.StandardKey.Open,
                                statusTip="Open an existing document", triggered=self.open_document)
        self.save_act = QAction("&Save", self, shortcut=QKeySequence.StandardKey.Save,
                                statusTip="Save the current document", triggered=self.save_document)
        self.save_as_act = QAction("Save &As...", self, shortcut=QKeySequence.StandardKey.SaveAs,
                                   statusTip="Save the current document with a new name",
                                   triggered=self.save_document_as)
        self.export_pdf_act = QAction("Export to &PDF...", self,
                                      statusTip="Export the current document as PDF",
                                      triggered=self.export_to_pdf)
        self.print_act = QAction("&Print...", self, shortcut=QKeySequence.StandardKey.Print,
                                 statusTip="Print the current document", triggered=self.print_document)
        self.print_preview_act = QAction("Print Pre&view...", self,
                                         statusTip="Preview how the document will look when printed",
                                         triggered=self.print_preview)
        self.close_tab_act = QAction("&Close Tab", self, shortcut="Ctrl+W",
                                     statusTip="Close the current tab", triggered=lambda: self.close_tab(self.tabs.currentIndex()))
        self.exit_act = QAction("E&xit", self, shortcut="Ctrl+Q",
                                statusTip="Exit ElbowOffice", triggered=self.close)

        # Edit actions
        self.undo_act = QAction("&Undo", self, shortcut=QKeySequence.StandardKey.Undo,
                                triggered=lambda: self.current_editor().undo() if self.current_editor() else None)
        self.redo_act = QAction("&Redo", self, shortcut=QKeySequence.StandardKey.Redo,
                                triggered=lambda: self.current_editor().redo() if self.current_editor() else None)
        self.cut_act = QAction("Cu&t", self, shortcut=QKeySequence.StandardKey.Cut,
                               triggered=lambda: self.current_editor().cut() if self.current_editor() else None)
        self.copy_act = QAction("&Copy", self, shortcut=QKeySequence.StandardKey.Copy,
                                triggered=lambda: self.current_editor().copy() if self.current_editor() else None)
        self.paste_act = QAction("&Paste", self, shortcut=QKeySequence.StandardKey.Paste,
                                 triggered=lambda: self.current_editor().paste() if self.current_editor() else None)
        self.select_all_act = QAction("Select &All", self, shortcut=QKeySequence.StandardKey.SelectAll,
                                      triggered=lambda: self.current_editor().selectAll() if self.current_editor() else None)
        self.find_replace_act = QAction("&Find && Replace...", self, shortcut="Ctrl+F",
                                        statusTip="Find and replace text", triggered=self.show_find_replace)

        # Format actions (will be connected to current editor)
        self.bold_act = QAction("&Bold", self, shortcut="Ctrl+B", checkable=True,
                                statusTip="Make the selected text bold")
        self.italic_act = QAction("&Italic", self, shortcut="Ctrl+I", checkable=True)
        self.underline_act = QAction("&Underline", self, shortcut="Ctrl+U", checkable=True)
        self.strike_act = QAction("&Strikethrough", self, checkable=True)

        self.align_left_act = QAction("Align &Left", self, checkable=True)
        self.align_center_act = QAction("Align &Center", self, checkable=True)
        self.align_right_act = QAction("Align &Right", self, checkable=True)
        self.align_justify_act = QAction("&Justify", self, checkable=True)

        self.bullet_list_act = QAction("• &Bullet List", self, checkable=True)
        self.number_list_act = QAction("1. &Numbered List", self, checkable=True)

        self.indent_more_act = QAction("Indent More", self, shortcut="Ctrl+]",
                                       statusTip="Increase paragraph indent")
        self.indent_less_act = QAction("Indent Less", self, shortcut="Ctrl+[",
                                       statusTip="Decrease paragraph indent")

        # Insert actions
        self.insert_image_act = QAction("Insert &Image...", self,
                                        statusTip="Insert an image from file",
                                        triggered=self.insert_image)
        self.insert_table_act = QAction("Insert &Table...", self,
                                        statusTip="Insert a table",
                                        triggered=self.insert_table)

        # View / Help
        self.zoom_in_act = QAction("Zoom &In", self, shortcut="Ctrl++",
                                   triggered=self.zoom_in)
        self.zoom_out_act = QAction("Zoom &Out", self, shortcut="Ctrl+-",
                                    triggered=self.zoom_out)
        self.reset_zoom_act = QAction("&Reset Zoom", self, shortcut="Ctrl+0",
                                      triggered=self.reset_zoom)

        self.about_act = QAction("&About ElbowOffice", self, triggered=self.show_about)

        # Connect format actions (these need special handling because they affect the current editor)
        self.bold_act.triggered.connect(self.toggle_bold)
        self.italic_act.triggered.connect(self.toggle_italic)
        self.underline_act.triggered.connect(self.toggle_underline)
        self.strike_act.triggered.connect(self.toggle_strike)

        self.align_left_act.triggered.connect(lambda: self.set_alignment(Qt.AlignmentFlag.AlignLeft))
        self.align_center_act.triggered.connect(lambda: self.set_alignment(Qt.AlignmentFlag.AlignCenter))
        self.align_right_act.triggered.connect(lambda: self.set_alignment(Qt.AlignmentFlag.AlignRight))
        self.align_justify_act.triggered.connect(lambda: self.set_alignment(Qt.AlignmentFlag.AlignJustify))

        self.bullet_list_act.triggered.connect(self.toggle_bullet_list)
        self.number_list_act.triggered.connect(self.toggle_number_list)

        self.indent_more_act.triggered.connect(lambda: self.change_indent(1))
        self.indent_less_act.triggered.connect(lambda: self.change_indent(-1))

    def _create_menus(self):
        menubar = self.menuBar()

        # File menu
        file_menu = menubar.addMenu("&File")
        file_menu.addAction(self.new_act)
        file_menu.addAction(self.open_act)
        file_menu.addSeparator()
        file_menu.addAction(self.save_act)
        file_menu.addAction(self.save_as_act)
        file_menu.addSeparator()
        file_menu.addAction(self.export_pdf_act)
        file_menu.addAction(self.print_act)
        file_menu.addAction(self.print_preview_act)
        file_menu.addSeparator()
        file_menu.addAction(self.close_tab_act)
        file_menu.addAction(self.exit_act)

        # Recent files submenu
        self.recent_menu = file_menu.addMenu("Open &Recent")
        self.recent_menu.aboutToShow.connect(self.update_recent_files_menu)

        # Edit menu
        edit_menu = menubar.addMenu("&Edit")
        edit_menu.addAction(self.undo_act)
        edit_menu.addAction(self.redo_act)
        edit_menu.addSeparator()
        edit_menu.addAction(self.cut_act)
        edit_menu.addAction(self.copy_act)
        edit_menu.addAction(self.paste_act)
        edit_menu.addSeparator()
        edit_menu.addAction(self.find_replace_act)
        edit_menu.addAction(self.select_all_act)

        # Format menu
        format_menu = menubar.addMenu("F&ormat")
        format_menu.addAction(self.bold_act)
        format_menu.addAction(self.italic_act)
        format_menu.addAction(self.underline_act)
        format_menu.addAction(self.strike_act)
        format_menu.addSeparator()

        align_menu = format_menu.addMenu("&Alignment")
        align_menu.addAction(self.align_left_act)
        align_menu.addAction(self.align_center_act)
        align_menu.addAction(self.align_right_act)
        align_menu.addAction(self.align_justify_act)

        list_menu = format_menu.addMenu("&Lists")
        list_menu.addAction(self.bullet_list_act)
        list_menu.addAction(self.number_list_act)

        format_menu.addAction(self.indent_more_act)
        format_menu.addAction(self.indent_less_act)

        # Insert menu
        insert_menu = menubar.addMenu("&Insert")
        insert_menu.addAction(self.insert_image_act)
        insert_menu.addAction(self.insert_table_act)

        # View menu
        view_menu = menubar.addMenu("&View")
        view_menu.addAction(self.zoom_in_act)
        view_menu.addAction(self.zoom_out_act)
        view_menu.addAction(self.reset_zoom_act)

        # Help menu
        help_menu = menubar.addMenu("&Help")
        help_menu.addAction(self.about_act)

    def _create_toolbars(self):
        # Main toolbar - File & Edit
        main_toolbar = QToolBar("Main", self)
        main_toolbar.setIconSize(QSize(18, 18))
        self.addToolBar(main_toolbar)

        main_toolbar.addAction(self.new_act)
        main_toolbar.addAction(self.open_act)
        main_toolbar.addAction(self.save_act)
        main_toolbar.addSeparator()
        main_toolbar.addAction(self.cut_act)
        main_toolbar.addAction(self.copy_act)
        main_toolbar.addAction(self.paste_act)
        main_toolbar.addSeparator()
        main_toolbar.addAction(self.print_act)

        # Formatting toolbar
        format_toolbar = QToolBar("Formatting", self)
        format_toolbar.setIconSize(QSize(18, 18))
        self.addToolBar(format_toolbar)

        # Font family
        self.font_combo = QFontComboBox()
        self.font_combo.setMaximumWidth(180)
        self.font_combo.currentFontChanged.connect(self.change_font_family)
        format_toolbar.addWidget(self.font_combo)

        # Font size
        self.size_combo = QComboBox()
        self.size_combo.addItems(["8", "9", "10", "11", "12", "14", "16", "18", "20", "22", "24", "28", "32", "36", "48", "56", "72"])
        self.size_combo.setCurrentText("11")
        self.size_combo.setEditable(True)
        self.size_combo.setMaximumWidth(60)
        self.size_combo.activated.connect(self.change_font_size)
        format_toolbar.addWidget(self.size_combo)

        format_toolbar.addSeparator()

        # Style toggles
        format_toolbar.addAction(self.bold_act)
        format_toolbar.addAction(self.italic_act)
        format_toolbar.addAction(self.underline_act)
        format_toolbar.addAction(self.strike_act)

        format_toolbar.addSeparator()

        # Colors
        self.text_color_act = QAction("Text Color", self, statusTip="Change text color")
        self.text_color_act.triggered.connect(self.change_text_color)
        format_toolbar.addAction(self.text_color_act)

        self.highlight_color_act = QAction("Highlight", self, statusTip="Change background highlight color")
        self.highlight_color_act.triggered.connect(self.change_highlight_color)
        format_toolbar.addAction(self.highlight_color_act)

        format_toolbar.addSeparator()

        # Alignment
        format_toolbar.addAction(self.align_left_act)
        format_toolbar.addAction(self.align_center_act)
        format_toolbar.addAction(self.align_right_act)
        format_toolbar.addAction(self.align_justify_act)

        format_toolbar.addSeparator()

        # Lists
        format_toolbar.addAction(self.bullet_list_act)
        format_toolbar.addAction(self.number_list_act)

        format_toolbar.addSeparator()
        format_toolbar.addAction(self.indent_less_act)
        format_toolbar.addAction(self.indent_more_act)

        # Insert toolbar
        insert_toolbar = QToolBar("Insert", self)
        insert_toolbar.setIconSize(QSize(18, 18))
        self.addToolBar(insert_toolbar)
        insert_toolbar.addAction(self.insert_image_act)
        insert_toolbar.addAction(self.insert_table_act)

    def _create_statusbar(self):
        self.status = QStatusBar()
        self.setStatusBar(self.status)

        # Left side info
        self.cursor_label = QLabel("Line 1, Col 1")
        self.status.addWidget(self.cursor_label)

        self.status.addPermanentWidget(QLabel("   "))

        # Right side - counts
        self.count_label = QLabel("Words: 0   Characters: 0")
        self.status.addPermanentWidget(self.count_label)

        # Modified indicator
        self.modified_label = QLabel("")
        self.status.addPermanentWidget(self.modified_label)

    # ---------------- DOCUMENT MANAGEMENT ----------------
    def current_tab(self) -> DocumentTab | None:
        widget = self.tabs.currentWidget()
        if isinstance(widget, DocumentTab):
            return widget
        return None

    def current_editor(self) -> QTextEdit | None:
        tab = self.current_tab()
        return tab.editor if tab else None

    def new_document(self):
        tab = DocumentTab()
        index = self.tabs.addTab(tab, tab.get_title())
        self.tabs.setCurrentIndex(index)

        # Connect signals for this editor
        editor = tab.editor
        editor.cursorPositionChanged.connect(self.update_cursor_position)
        editor.textChanged.connect(self.update_counts_and_title)
        editor.currentCharFormatChanged.connect(self.update_format_toolbar)
        editor.selectionChanged.connect(self.update_format_toolbar)

        # Set default font
        editor.setFont(QFont("Calibri", 11))

        self.update_counts_and_title()
        self.update_format_toolbar()

        # Give focus to the editor
        editor.setFocus()

        # First-time friendly content
        if self.tabs.count() == 1 and not tab.file_path:
            welcome = """<h2>Welcome to ElbowOffice!</h2>
<p>This is a simple but capable word processor inspired by OpenOffice Writer.</p>
<p><b>Try these things:</b></p>
<ul>
  <li>Select text and use the <b>Bold</b>, <i>Italic</i>, and <u>Underline</u> buttons</li>
  <li>Change the font family and size using the dropdowns</li>
  <li>Click the color buttons to change text or highlight color</li>
  <li>Use the alignment and list buttons on the toolbar</li>
  <li>Go to <b>Insert → Image</b> or <b>Insert → Table</b></li>
</ul>
<p>Everything you type here is fully formatted rich text.</p>
<p>Happy writing!</p>"""
            editor.setHtml(welcome)
            tab.set_modified(False)  # Don't count the welcome as a modification

    def close_tab(self, index: int):
        if index < 0:
            return

        tab = self.tabs.widget(index)
        if not isinstance(tab, DocumentTab):
            self.tabs.removeTab(index)
            return

        if tab.is_modified():
            reply = QMessageBox.question(
                self, "Unsaved Changes",
                f"Do you want to save changes to '{tab.get_title()}'?",
                QMessageBox.StandardButton.Save | QMessageBox.StandardButton.Discard | QMessageBox.StandardButton.Cancel
            )
            if reply == QMessageBox.StandardButton.Save:
                if not self.save_document_for_tab(tab):
                    return  # User cancelled save
            elif reply == QMessageBox.StandardButton.Cancel:
                return

        self.tabs.removeTab(index)

        # If no tabs left, create a new one
        if self.tabs.count() == 0:
            self.new_document()

    def on_tab_changed(self, index):
        self.update_counts_and_title()
        self.update_format_toolbar()
        self.update_window_title()

        # Reconnect cursor signal if needed
        tab = self.current_tab()
        if tab:
            try:
                tab.editor.cursorPositionChanged.disconnect(self.update_cursor_position)
            except:
                pass
            tab.editor.cursorPositionChanged.connect(self.update_cursor_position)

    def update_window_title(self):
        tab = self.current_tab()
        if tab:
            title = f"{tab.get_title()} - ElbowOffice"
            if tab.is_modified():
                title = f"*{tab.get_title()} - ElbowOffice"
            self.setWindowTitle(title)

    def update_counts_and_title(self):
        tab = self.current_tab()
        if not tab:
            return

        words = tab.word_count()
        chars = tab.char_count()
        self.count_label.setText(f"Words: {words}   Characters: {chars}")

        # Update tab text + window title
        title = tab.get_title()
        if tab.is_modified():
            title = f"*{title}"
            self.modified_label.setText("Modified")
        else:
            self.modified_label.setText("")

        index = self.tabs.currentIndex()
        self.tabs.setTabText(index, title)
        self.update_window_title()

    def update_cursor_position(self):
        editor = self.current_editor()
        if not editor:
            return
        cursor = editor.textCursor()
        line = cursor.blockNumber() + 1
        col = cursor.columnNumber() + 1
        self.cursor_label.setText(f"Line {line}, Col {col}")

    # ---------------- FORMATTING ----------------
    def update_format_toolbar(self):
        """Sync toolbar buttons and combos with current cursor format."""
        editor = self.current_editor()
        if not editor:
            return

        cursor = editor.textCursor()
        char_format = cursor.charFormat()
        block_format = cursor.blockFormat()

        # Font family
        font = char_format.font()
        self.font_combo.blockSignals(True)
        self.font_combo.setCurrentFont(font)
        self.font_combo.blockSignals(False)

        # Font size
        size = int(font.pointSize())
        if size > 0:
            self.size_combo.blockSignals(True)
            if str(size) in [self.size_combo.itemText(i) for i in range(self.size_combo.count())]:
                self.size_combo.setCurrentText(str(size))
            else:
                self.size_combo.setCurrentText(str(size))
            self.size_combo.blockSignals(False)

        # Style toggles
        self.bold_act.setChecked(char_format.fontWeight() == QFont.Weight.Bold)
        self.italic_act.setChecked(char_format.fontItalic())
        self.underline_act.setChecked(char_format.fontUnderline())
        self.strike_act.setChecked(char_format.fontStrikeOut())

        # Alignment
        align = block_format.alignment()
        self.align_left_act.setChecked(align == Qt.AlignmentFlag.AlignLeft)
        self.align_center_act.setChecked(align == Qt.AlignmentFlag.AlignCenter)
        self.align_right_act.setChecked(align == Qt.AlignmentFlag.AlignRight)
        self.align_justify_act.setChecked(align == Qt.AlignmentFlag.AlignJustify)

        # Lists
        list_format = cursor.currentList()
        self.bullet_list_act.setChecked(bool(list_format and list_format.style() in (
            QTextListFormat.Style.ListDisc,
            QTextListFormat.Style.ListCircle,
            QTextListFormat.Style.ListSquare
        )))
        self.number_list_act.setChecked(bool(list_format and list_format.style() in (
            QTextListFormat.Style.ListDecimal,
            QTextListFormat.Style.ListLowerAlpha,
            QTextListFormat.Style.ListUpperAlpha
        )))

    def change_font_family(self, font: QFont):
        editor = self.current_editor()
        if not editor:
            return
        fmt = QTextCharFormat()
        fmt.setFontFamilies([font.family()])
        self.merge_format(fmt)

    def change_font_size(self):
        editor = self.current_editor()
        if not editor:
            return
        try:
            size = int(self.size_combo.currentText())
        except ValueError:
            return
        fmt = QTextCharFormat()
        fmt.setFontPointSize(size)
        self.merge_format(fmt)

    def merge_format(self, fmt: QTextCharFormat):
        editor = self.current_editor()
        if not editor:
            return
        cursor = editor.textCursor()
        if not cursor.hasSelection():
            cursor.select(QTextCursor.SelectionType.WordUnderCursor)
        cursor.mergeCharFormat(fmt)
        editor.mergeCurrentCharFormat(fmt)

    def toggle_bold(self):
        editor = self.current_editor()
        if not editor:
            return
        fmt = QTextCharFormat()
        weight = QFont.Weight.Bold if self.bold_act.isChecked() else QFont.Weight.Normal
        fmt.setFontWeight(weight)
        self.merge_format(fmt)

    def toggle_italic(self):
        editor = self.current_editor()
        if not editor:
            return
        fmt = QTextCharFormat()
        fmt.setFontItalic(self.italic_act.isChecked())
        self.merge_format(fmt)

    def toggle_underline(self):
        editor = self.current_editor()
        if not editor:
            return
        fmt = QTextCharFormat()
        fmt.setFontUnderline(self.underline_act.isChecked())
        self.merge_format(fmt)

    def toggle_strike(self):
        editor = self.current_editor()
        if not editor:
            return
        fmt = QTextCharFormat()
        fmt.setFontStrikeOut(self.strike_act.isChecked())
        self.merge_format(fmt)

    def set_alignment(self, alignment):
        editor = self.current_editor()
        if not editor:
            return
        cursor = editor.textCursor()
        block_format = cursor.blockFormat()
        block_format.setAlignment(alignment)
        cursor.setBlockFormat(block_format)
        editor.setTextCursor(cursor)

        # Update exclusive check state
        self.align_left_act.setChecked(alignment == Qt.AlignmentFlag.AlignLeft)
        self.align_center_act.setChecked(alignment == Qt.AlignmentFlag.AlignCenter)
        self.align_right_act.setChecked(alignment == Qt.AlignmentFlag.AlignRight)
        self.align_justify_act.setChecked(alignment == Qt.AlignmentFlag.AlignJustify)

    def toggle_bullet_list(self):
        self._toggle_list(QTextListFormat.Style.ListDisc)

    def toggle_number_list(self):
        self._toggle_list(QTextListFormat.Style.ListDecimal)

    def _toggle_list(self, style):
        editor = self.current_editor()
        if not editor:
            return

        cursor = editor.textCursor()
        current_list = cursor.currentList()

        if current_list:
            # Remove list
            cursor.createList(QTextListFormat.Style.ListStyleUndefined)
        else:
            # Create list
            list_format = QTextListFormat()
            list_format.setStyle(style)
            cursor.createList(list_format)

        self.update_format_toolbar()

    def change_indent(self, delta: int):
        editor = self.current_editor()
        if not editor:
            return
        cursor = editor.textCursor()
        block_format = cursor.blockFormat()
        indent = block_format.indent() + delta
        indent = max(0, min(12, indent))  # reasonable limits
        block_format.setIndent(indent)
        cursor.setBlockFormat(block_format)
        editor.setTextCursor(cursor)

    def change_text_color(self):
        editor = self.current_editor()
        if not editor:
            return
        color = QColorDialog.getColor(editor.textColor(), self, "Select Text Color")
        if color.isValid():
            fmt = QTextCharFormat()
            fmt.setForeground(color)
            self.merge_format(fmt)

    def change_highlight_color(self):
        editor = self.current_editor()
        if not editor:
            return
        color = QColorDialog.getColor(Qt.GlobalColor.yellow, self, "Select Highlight Color")
        if color.isValid():
            fmt = QTextCharFormat()
            fmt.setBackground(color)
            self.merge_format(fmt)

    # ---------------- FILE OPERATIONS ----------------
    def open_document(self):
        file_filter = "All Supported (*.html *.htm *.txt *.docx);;HTML Files (*.html *.htm);;Text Files (*.txt)"
        if DOCX_AVAILABLE:
            file_filter = "All Supported (*.html *.htm *.txt *.docx);;HTML Files (*.html *.htm);;Text Files (*.txt);;Word Documents (*.docx)"

        path, _ = QFileDialog.getOpenFileName(
            self, "Open Document", get_desktop_path(), file_filter
        )
        if not path:
            return

        self.load_file(path)

    def load_file(self, path: str):
        if not os.path.exists(path):
            QMessageBox.warning(self, "File Not Found", f"Could not find:\n{path}")
            return

        tab = DocumentTab()
        editor = tab.editor

        ext = os.path.splitext(path)[1].lower()

        try:
            if ext in (".html", ".htm"):
                with open(path, "r", encoding="utf-8", errors="replace") as f:
                    html = f.read()
                editor.setHtml(html)
            elif ext == ".docx" and DOCX_AVAILABLE:
                self._load_docx(tab, path)
            else:
                # Plain text
                with open(path, "r", encoding="utf-8", errors="replace") as f:
                    text = f.read()
                editor.setPlainText(text)

            tab.file_path = path
            tab.set_modified(False)

            index = self.tabs.addTab(tab, tab.get_title())
            self.tabs.setCurrentIndex(index)

            # Wire up signals
            editor.cursorPositionChanged.connect(self.update_cursor_position)
            editor.textChanged.connect(self.update_counts_and_title)
            editor.currentCharFormatChanged.connect(self.update_format_toolbar)
            editor.selectionChanged.connect(self.update_format_toolbar)

            self.add_to_recent_files(path)
            self.update_counts_and_title()
            self.statusBar().showMessage(f"Opened: {os.path.basename(path)}", 3000)

        except Exception as e:
            QMessageBox.critical(self, "Error Opening File", f"Could not open file:\n{str(e)}")

    def _load_docx(self, tab: DocumentTab, path: str):
        """Convert .docx to HTML for the editor (basic but usable)."""
        doc = DocxDocument(path)
        html_parts = ["<html><body>"]

        for para in doc.paragraphs:
            style = ""
            if para.style and "Heading" in para.style.name:
                level = para.style.name.split()[-1] if para.style.name.split()[-1].isdigit() else "1"
                tag = f"h{level}"
            else:
                tag = "p"

            text = ""
            for run in para.runs:
                rtext = run.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                if run.bold:
                    rtext = f"<b>{rtext}</b>"
                if run.italic:
                    rtext = f"<i>{rtext}</i>"
                if run.underline:
                    rtext = f"<u>{rtext}</u>"
                text += rtext

            align = ""
            if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                align = ' style="text-align:center"'
            elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                align = ' style="text-align:right"'

            html_parts.append(f"<{tag}{align}>{text}</{tag}>")

        html_parts.append("</body></html>")
        tab.editor.setHtml("".join(html_parts))

    def save_document(self):
        tab = self.current_tab()
        if not tab:
            return False
        if tab.file_path:
            return self._save_tab_to_path(tab, tab.file_path)
        else:
            return self.save_document_as()

    def save_document_as(self):
        tab = self.current_tab()
        if not tab:
            return False
        return self.save_document_for_tab(tab)

    def save_document_for_tab(self, tab: DocumentTab) -> bool:
        default_name = tab.get_title()
        if not default_name.endswith((".html", ".txt", ".docx")):
            default_name += ".html"

        filters = "HTML Document (*.html);;Plain Text (*.txt)"
        if DOCX_AVAILABLE:
            filters = "HTML Document (*.html);;Word Document (*.docx);;Plain Text (*.txt)"

        path, selected_filter = QFileDialog.getSaveFileName(
            self, "Save Document", os.path.join(get_desktop_path(), default_name),
            filters
        )
        if not path:
            return False

        # If user chose .docx in the dialog
        if path.lower().endswith(".docx") and DOCX_AVAILABLE:
            return self._save_as_docx(tab, path)

        return self._save_tab_to_path(tab, path)

    def _save_tab_to_path(self, tab: DocumentTab, path: str) -> bool:
        ext = os.path.splitext(path)[1].lower()
        try:
            if ext == ".txt":
                with open(path, "w", encoding="utf-8") as f:
                    f.write(tab.editor.toPlainText())
            else:
                # Default to HTML (best fidelity for rich text)
                if not path.lower().endswith((".html", ".htm")):
                    path += ".html"
                with open(path, "w", encoding="utf-8") as f:
                    f.write(tab.editor.toHtml())

            tab.file_path = path
            tab.set_modified(False)
            self.add_to_recent_files(path)

            # Refresh UI
            idx = self.tabs.indexOf(tab)
            if idx >= 0:
                self.tabs.setTabText(idx, tab.get_title())
            self.update_counts_and_title()
            self.statusBar().showMessage(f"Saved: {os.path.basename(path)}", 3000)
            return True

        except Exception as e:
            QMessageBox.critical(self, "Save Error", f"Could not save file:\n{str(e)}")
            return False

    def _save_as_docx(self, tab: DocumentTab, path: str) -> bool:
        """Export the current document to .docx format using python-docx (basic formatting preserved)."""
        if not DOCX_AVAILABLE:
            QMessageBox.warning(self, "DOCX Unavailable",
                                "python-docx is not installed.\n\nInstall it with:\n    pip install python-docx")
            return False

        try:
            doc = DocxDocument()

            # Walk the QTextDocument
            document = tab.editor.document()
            block = document.begin()

            while block.isValid():
                para = doc.add_paragraph()
                para_format = para.paragraph_format

                # Alignment
                align = block.blockFormat().alignment()
                if align == Qt.AlignmentFlag.AlignCenter:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == Qt.AlignmentFlag.AlignRight:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif align == Qt.AlignmentFlag.AlignJustify:
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Process fragments (runs)
                it = block.begin()
                while not it.atEnd():
                    fragment = it.fragment()
                    if fragment.isValid():
                        char_format = fragment.charFormat()
                        text = fragment.text()

                        run = para.add_run(text)

                        # Basic formatting
                        if char_format.fontWeight() == QFont.Weight.Bold:
                            run.bold = True
                        if char_format.fontItalic():
                            run.italic = True
                        if char_format.fontUnderline():
                            run.underline = True
                        if char_format.fontStrikeOut():
                            run.font.strike = True

                        # Color
                        if char_format.foreground().color().isValid():
                            c = char_format.foreground().color()
                            run.font.color.rgb = RGBColor(c.red(), c.green(), c.blue())

                        # Size
                        if char_format.fontPointSize() > 0:
                            run.font.size = Pt(char_format.fontPointSize())

                    it += 1

                block = block.next()

            doc.save(path)

            tab.file_path = path
            tab.set_modified(False)
            self.add_to_recent_files(path)

            idx = self.tabs.indexOf(tab)
            if idx >= 0:
                self.tabs.setTabText(idx, tab.get_title())
            self.update_counts_and_title()
            self.statusBar().showMessage(f"Saved DOCX: {os.path.basename(path)}", 3000)
            return True

        except Exception as e:
            QMessageBox.critical(self, "DOCX Export Error", f"Failed to save .docx file:\n{str(e)}")
            return False

    def export_to_pdf(self):
        tab = self.current_tab()
        if not tab:
            return

        path, _ = QFileDialog.getSaveFileName(
            self, "Export to PDF", os.path.join(get_desktop_path(), tab.get_title() + ".pdf"),
            "PDF Files (*.pdf)"
        )
        if not path:
            return

        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        printer.setOutputFileName(path)
        printer.setPageMargins(QPageLayout.Margins(15, 15, 15, 15, QPageLayout.Unit.Millimeter))

        doc = tab.editor.document().clone()
        doc.print(printer)

        self.statusBar().showMessage(f"Exported PDF: {os.path.basename(path)}", 4000)

    def print_document(self):
        editor = self.current_editor()
        if not editor:
            return

        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        dialog = QPrintDialog(printer, self)
        if dialog.exec() == QPrintDialog.DialogCode.Accepted:
            editor.print(printer)

    def print_preview(self):
        editor = self.current_editor()
        if not editor:
            return

        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        preview = QPrintPreviewDialog(printer, self)
        preview.paintRequested.connect(editor.print)
        preview.exec()

    # ---------------- INSERT FEATURES ----------------
    def insert_image(self):
        editor = self.current_editor()
        if not editor:
            return

        path, _ = QFileDialog.getOpenFileName(
            self, "Insert Image", get_desktop_path(),
            "Images (*.png *.jpg *.jpeg *.gif *.bmp *.webp)"
        )
        if not path:
            return

        cursor = editor.textCursor()
        image_format = QTextImageFormat()
        image_format.setName(path)
        # Reasonable default width
        image_format.setWidth(320)
        cursor.insertImage(image_format)

    def insert_table(self):
        editor = self.current_editor()
        if not editor:
            return

        rows, ok = QInputDialog.getInt(self, "Insert Table", "Number of rows:", 3, 1, 30)
        if not ok:
            return
        cols, ok = QInputDialog.getInt(self, "Insert Table", "Number of columns:", 3, 1, 10)
        if not ok:
            return

        cursor = editor.textCursor()
        table = cursor.insertTable(rows, cols)

        # Style the table a little (borders via html is limited in QTextEdit, but we can set cell padding)
        for r in range(rows):
            for c in range(cols):
                cell = table.cellAt(r, c)
                cell.setFormat(cell.format())  # placeholder for future styling

    # ---------------- FIND / REPLACE ----------------
    def show_find_replace(self):
        if not self.current_editor():
            return

        if self.current_find_dialog is None:
            self.current_find_dialog = FindReplaceDialog(self)
            self.current_find_dialog.find_next_btn.clicked.connect(self.find_next)
            self.current_find_dialog.replace_btn.clicked.connect(self.replace_once)
            self.current_find_dialog.replace_all_btn.clicked.connect(self.replace_all)

        self.current_find_dialog.show()
        self.current_find_dialog.raise_()
        self.current_find_dialog.find_edit.setFocus()
        self.current_find_dialog.find_edit.selectAll()

    def _get_find_flags(self):
        flags = QTextDocument.FindFlag(0)
        dlg = self.current_find_dialog
        if dlg:
            if dlg.case_sensitive.isChecked():
                flags |= QTextDocument.FindFlag.FindCaseSensitively
            if dlg.whole_word.isChecked():
                flags |= QTextDocument.FindFlag.FindWholeWords
        return flags

    def find_next(self):
        editor = self.current_editor()
        if not editor or not self.current_find_dialog:
            return
        text = self.current_find_dialog.find_edit.text()
        if not text:
            return

        found = editor.find(text, self._get_find_flags())
        if not found:
            # wrap around
            cursor = editor.textCursor()
            cursor.movePosition(QTextCursor.MoveOperation.Start)
            editor.setTextCursor(cursor)
            editor.find(text, self._get_find_flags())

    def replace_once(self):
        editor = self.current_editor()
        if not editor or not self.current_find_dialog:
            return
        find_text = self.current_find_dialog.find_edit.text()
        replace_text = self.current_find_dialog.replace_edit.text()
        if not find_text:
            return

        cursor = editor.textCursor()
        if cursor.hasSelection() and cursor.selectedText().lower() == find_text.lower():
            cursor.insertText(replace_text)
        self.find_next()

    def replace_all(self):
        editor = self.current_editor()
        if not editor or not self.current_find_dialog:
            return
        find_text = self.current_find_dialog.find_edit.text()
        replace_text = self.current_find_dialog.replace_edit.text()
        if not find_text:
            return

        flags = self._get_find_flags()
        cursor = editor.textCursor()
        cursor.beginEditBlock()

        count = 0
        # Start from beginning
        cursor.movePosition(QTextCursor.MoveOperation.Start)
        editor.setTextCursor(cursor)

        while True:
            if not editor.find(find_text, flags):
                break
            c = editor.textCursor()
            c.insertText(replace_text)
            count += 1

        cursor.endEditBlock()
        QMessageBox.information(self, "Replace All", f"Replaced {count} occurrence(s).")

    # ---------------- ZOOM ----------------
    def zoom_in(self):
        editor = self.current_editor()
        if editor:
            editor.zoomIn(1)

    def zoom_out(self):
        editor = self.current_editor()
        if editor:
            editor.zoomOut(1)

    def reset_zoom(self):
        editor = self.current_editor()
        if editor:
            # Reset to default font size
            font = editor.font()
            font.setPointSize(11)
            editor.setFont(font)
            # Also reset any zoom factor applied by zoomIn/Out
            editor.setFontPointSize(11)

    # ---------------- RECENT FILES ----------------
    def add_to_recent_files(self, path: str):
        if path in self.recent_files:
            self.recent_files.remove(path)
        self.recent_files.insert(0, path)
        self.recent_files = self.recent_files[:10]  # keep last 10
        self.settings.setValue("recent_files", self.recent_files)
        self.update_recent_files_menu()

    def update_recent_files_menu(self):
        self.recent_menu.clear()
        if not self.recent_files:
            action = QAction("(No recent files)", self)
            action.setEnabled(False)
            self.recent_menu.addAction(action)
            return

        for path in self.recent_files:
            action = QAction(os.path.basename(path), self)
            action.setStatusTip(path)
            action.triggered.connect(lambda checked, p=path: self.load_file(p))
            self.recent_menu.addAction(action)

        self.recent_menu.addSeparator()
        clear_act = QAction("Clear Recent Files", self)
        clear_act.triggered.connect(self.clear_recent_files)
        self.recent_menu.addAction(clear_act)

    def clear_recent_files(self):
        self.recent_files = []
        self.settings.setValue("recent_files", [])
        self.update_recent_files_menu()

    # ---------------- ABOUT & CLOSE ----------------
    def show_about(self):
        about_text = f"""
<b>ElbowOffice</b><br>
A simple but capable word processor written in Python + PyQt6<br><br>
Inspired by OpenOffice Writer and LibreOffice Writer.<br><br>
<b>Features:</b><br>
• Rich text formatting<br>
• Multiple tabs<br>
• Images &amp; tables<br>
• HTML, TXT, and DOCX (optional) support<br>
• PDF export &amp; printing<br><br>
<b>Dependencies:</b><br>
PyQt6 (required)<br>
python-docx (optional, for .docx)<br><br>
Version 1.0 — Made for learning and everyday writing. (ElbowOffice)
"""
        QMessageBox.about(self, "About ElbowOffice", about_text)

    def closeEvent(self, event):
        # Check all tabs for unsaved changes
        unsaved = []
        for i in range(self.tabs.count()):
            tab = self.tabs.widget(i)
            if isinstance(tab, DocumentTab) and tab.is_modified():
                unsaved.append(tab.get_title())

        if unsaved:
            msg = "You have unsaved changes in:\n\n" + "\n".join(f"• {t}" for t in unsaved)
            reply = QMessageBox.question(
                self, "Unsaved Changes",
                msg + "\n\nDo you really want to quit?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.No:
                event.ignore()
                return

        # Save settings
        self.settings.setValue("geometry", self.saveGeometry())
        self.settings.setValue("recent_files", self.recent_files)

        event.accept()


# ====================== MAIN ENTRY ======================

def main():
    app = QApplication(sys.argv)
    app.setApplicationName("ElbowOffice")
    app.setOrganizationName("ElbowOffice")

    # Nice default style
    app.setStyle("Fusion")

    window = ElbowOffice()
    window.show()

    sys.exit(app.exec())


if __name__ == "__main__":
    main()
